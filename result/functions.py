import numpy as np
import pandas as pd
import os
import xml.dom.minidom
import re
from result.models import *
from django.core.exceptions import ObjectDoesNotExist
from django.shortcuts import get_object_or_404
from django.conf import settings
import json
import matplotlib
matplotlib.use("agg")
import matplotlib.pyplot as plt
from mailmerge import MailMerge
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from datetime import datetime
import math
from decimal import Decimal
from django.contrib import messages
import random

cycles = np.arange(1, 41)
def no_accent_vietnamese(s):    
    s = re.sub(r'[àáạảãâầấậẩẫăằắặẳẵ]', 'a', s)
    s = re.sub(r'[ÀÁẠẢÃĂẰẮẶẲẴÂẦẤẬẨẪ]', 'A', s)
    s = re.sub(r'[[èéẹẻẽêềếệểễ]', 'e', s)
    s = re.sub(r'[ÈÉẸẺẼÊỀẾỆỂỄ]', 'E', s)
    s = re.sub(r'[òóọỏõôồốộổỗơờớợởỡ]', 'o', s)
    s = re.sub(r'[ÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠ]', 'O', s)
    s = re.sub(r'[ìíịỉĩ]', 'i', s)
    s = re.sub(r'[ÌÍỊỈĨ]', 'I', s)
    s = re.sub(r'[ùúụủũưừứựửữ]', 'u', s)
    s = re.sub(r'[ƯỪỨỰỬỮÙÚỤỦŨ]', 'U', s)
    s = re.sub(r'[ỳýỵỷỹ]', 'y', s)
    s = re.sub(r'[ỲÝỴỶỸ]', 'Y', s)
    s = re.sub(r'[Đ]', 'D', s)
    s = re.sub(r'[đ]', 'd', s)
    return s

#HBV FUNCTIONS
def create_HBV_image(pk = None):
    try:
        cycles = np.arange(1, 41)
        obj = get_object_or_404(HBVSample, pk = pk)
        completed = ""
        incompleted = ""
        #try:
        std = obj.std_curve
        standards = json.loads(std.stds)
        quant_curve = json.loads(std.std_curve)
        curves = json.loads(obj.curves) 

        #colors and makers for ax1
        color = ['#0084E8','#0C40FF','#2300E8']
        marker = ['o', '<','x']
        fig, (ax1,ax2) = plt.subplots(1,2)
        fig.set_size_inches(10, 4)
        #plot duong chay
        for no, name in list(enumerate(standards)):    
            ax1.plot(cycles, pd.DataFrame(standards[name]), label = 'Standard ' + str(name), c = color[no], marker = marker[no], markersize = 3, linestyle = '-', linewidth = 1)
        ax1.plot(cycles, pd.DataFrame(curves["green"]), label = obj.lab_id, c = 'red', linewidth = 2.3)
        ax1.plot(cycles, pd.DataFrame(curves['yellow']), label = 'Chứng nội', c = 'green', alpha = 0.7, marker = '+', markersize = 3)
        ax1.legend()
        ax1.set_xlabel('Cycles')
        ax1.set_ylabel('Normalized Fluorescence')    
        ax1.set_title('Amplification Plot')
        ax1.grid(linestyle = ':', alpha = 0.5)
        ax1.hlines(0.02, xmin = 0, xmax = 40, colors='grey', linestyles='dotted', linewidth = 0.7)
        ax1.text(0, 0.13, 'Threshold', fontsize = 9, color='grey')
        #plot duong chuan
        line_x = pd.DataFrame([1, 8])
        line_y = pd.DataFrame([std.slope + std.B, std.slope * 8 + std.B])
        ax2.plot(line_x, line_y, c = 'black')
        for k,v in quant_curve.items():
            ax2.scatter(float(k),v, c = 'black',s = 10, marker = 's')
        ax2.set_title("Standard Curve")
        ax2.set_xlabel("Log concentration per reaction")
        ax2.set_ylabel('C$_t$')
        
        if obj.ct != None:
            ax2.scatter((obj.ct - std.B)/std.slope, obj.ct, label = obj.lab_id)
        
        ax2.legend()
        ax2.grid(linestyle = ':', alpha = 0.5)
        textstr = '\n'.join((
        r'Efficiency = %.2f %%' % (std.eff, ),
        r'$R^2$ = %.2f' % (std.r_2, ),
        r'Slope = %.2f' % (std.slope, )))
        fig.tight_layout(h_pad = 3)    
        props = dict(boxstyle='round', alpha=0)
        ax2.text(0.05, 0.05, textstr, transform=ax2.transAxes,
            verticalalignment='bottom', bbox=props)
        name = "static/run_img/"+ obj.lab_id +'_amp.png'
        fig.savefig(name, type = 'png', bbox_inches = 'tight', dpi= 100)
        plt.close("all")
        completed = obj.lab_id
        
    
    except TypeError:
        incompleted = "Chưa nhập file chạy " + obj.lab_id
    
    return completed, incompleted

def process_HBV_runfile(green_file, yellow_file):
    completed = []
    not_complete = []
    
    try:
        file_1 = xml.dom.minidom.parse(green_file)
        root_1 = file_1.documentElement
        data_1 = {}
        series = root_1.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Hh][Bb][vV]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_1[name] = array

        file_2 = xml.dom.minidom.parse(yellow_file)
        root_2 = file_2.documentElement
        data_2 = {}
        series = root_2.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Hh][Bb][vV]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_2[name] = array
    except AttributeError:
        not_complete.append('File chạy không chứa mẫu tên HBV')

    else:
        for k, name in list(enumerate(data_1)):
            try:
                pt = HBVSample.objects.filter(lab_id=name).order_by('added').first()
                js_data = {"green": data_1[name], "yellow": data_2[name]}
                js = json.dumps(js_data)
                pt.curves = js
                pt.save()

                completed.append(name)
            except(HBVSample.DoesNotExist, ObjectDoesNotExist):

                not_complete.append(name)
            except KeyError:
                not_complete.append('Hai files chạy không khớp nhau')
    for id_ in completed:
        obj = HBVSample.objects.get(lab_id = id_)
        print(obj.pk)
        c,i = create_HBV_image(obj.pk)
    
    return completed, not_complete

def create_HBV_report(pk = None):
    obj = HBVSample.objects.get(pk = pk)
    completed = ""
    incompleted = ""
    try:
        #desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
        today_folder =  "D:/Result/" + datetime.now().strftime("%m%d")
        if not os.path.exists(today_folder):
            os.makedirs(today_folder, exist_ok = True)
        try:
            print(obj.sid)
            sid_short = re.search(r"\d+-(\d{4})", obj.sid).groups()[0]

        except (AttributeError, TypeError):
            sid_short = ""
        file_name = today_folder +"/" + sid_short + '-' + no_accent_vietnamese(obj.name.replace(" ",'')) + '-' +obj.lab_id+'.docx'
        pic_file = settings.STATIC_DIR + "/run_img/" + obj.lab_id + '_amp.png'
        if obj.sid == None:
            sid = ""
        else:
            sid = str(obj.sid)
        
        if obj.age != None:
            age = str(obj.age)
        else:
            age = ""
        #start merging
        f_1 = MailMerge(settings.STATIC_DIR + '/word_template/HBV FORM.docx')
        f_1.merge(
            sid = sid, 
            lab_id = str(obj.lab_id), 
            sex = obj.sex, 
            doctor = obj.doctor, 
            addres = obj.address, 
            age = age, 
            name = obj.name, 
            date_receive = obj.added.strftime('%d/%m/%Y'), 
            date_finished = obj.modified.strftime('%d/%m/%Y-%H:%M'), 
            d= datetime.now().strftime("%d"), 
            m= datetime.now().strftime('%m'), 
            y= datetime.now().strftime('%Y'))      
        if obj.result == "DƯỚI NGƯỠNG PHÁT HIỆN":
            f_1.merge(iu = "DƯỚI NGƯỠNG PHÁT HIỆN")
        elif obj.copies != None:        
            try:            
                f_1.merge(
                copies =  "{:.1E}".format(Decimal(obj.copies*obj.std_curve.factor)),
                iu = "{:.1E}".format(Decimal(obj.copies*obj.std_curve.factor/5)),
                log = "{:.1f}".format(math.log10(obj.copies*obj.std_curve.factor)))
            except:
                print('ERROR')

        f_1.write(file_name)

        doc = Document(file_name)
        for paragraph in doc.paragraphs:
            if r"None" in paragraph.text:
                print('YES NONE IS HERE')
                paragraph.text = paragraph.text.strip().replace("None", "")
            if "[image]" in paragraph.text:
                paragraph.text = paragraph.text.strip().replace("[image]", "")

                run = paragraph.add_run()
                run.add_picture(pic_file, width=Cm(15))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(file_name)
        completed = obj.lab_id + " vào " + today_folder
    except FileNotFoundError:
        incompleted = "Chưa có file hình " + obj.lab_id
    except:
        incompleted = obj.lab_id
    return completed, incompleted

#HCV FUNCTIONS
def create_HCV_image(pk = None):
    print('IM HERE')
    cycles = np.arange(1, 41)
    obj = get_object_or_404(HCVSample, pk = pk)
    completed = ""
    incompleted = ""
    #try:
    std = obj.std_curve
    standards = json.loads(std.stds)
    quant_curve = json.loads(std.std_curve)
    curves = json.loads(obj.curves) 
    print(type(standards))
    print(standards)
    #colors and makers for ax1
    color = ['#0084E8','#0C40FF','#2300E8']
    marker = ['o', '<','x']
    fig, (ax1,ax2) = plt.subplots(1,2)
    fig.set_size_inches(10, 4)
    #plot duong chay
    for no, name in list(enumerate(standards)):    
        ax1.plot(cycles, pd.DataFrame(standards[name]), label = 'Standard ' + str(name), c = color[no], marker = marker[no], markersize = 3, linestyle = '-', linewidth = 1)
    ax1.plot(cycles, pd.DataFrame(curves["green"]), label = obj.lab_id, c = 'red', linewidth = 2.3)
    ax1.plot(cycles, pd.DataFrame(curves['yellow']), label = 'Chứng nội', c = 'green', alpha = 0.7, marker = '+', markersize = 3)
    ax1.legend()
    ax1.set_xlabel('Cycles')
    ax1.set_ylabel('Normalized Fluorescence')    
    ax1.set_title('Amplification Plot')
    ax1.grid(linestyle = ':', alpha = 0.5)
    ax1.hlines(0.02, xmin = 0, xmax = 40, colors='grey', linestyles='dotted', linewidth = 0.7)
    ax1.text(0, 0.13, 'Threshold', fontsize = 9, color='grey')
    #plot duong chuan
    line_x = pd.DataFrame([1, 8])
    line_y = pd.DataFrame([std.slope + std.B, std.slope * 8 + std.B])
    ax2.plot(line_x, line_y, c = 'black')
    for k,v in quant_curve.items():
        ax2.scatter(float(k),v, c = 'black',s = 10, marker = 's')
    ax2.set_title("Standard Curve")
    ax2.set_xlabel("Log concentration per reaction")
    ax2.set_ylabel('C$_t$')
    
    if obj.ct != None:
        ax2.scatter((obj.ct - std.B)/std.slope, obj.ct, label = obj.lab_id)
    
    ax2.legend()
    ax2.grid(linestyle = ':', alpha = 0.5)
    textstr = '\n'.join((
    r'Efficiency = %.2f %%' % (std.eff, ),
    r'$R^2$ = %.2f' % (std.r_2, ),
    r'Slope = %.2f' % (std.slope, )))
    fig.tight_layout(h_pad = 3)    
    props = dict(boxstyle='round', alpha=0)
    ax2.text(0.05, 0.05, textstr, transform=ax2.transAxes,
        verticalalignment='bottom', bbox=props)
    name = "static/run_img/"+ obj.lab_id +'_amp.png'
    fig.savefig(name, type = 'png', bbox_inches = 'tight', dpi= 100)
    completed = obj.lab_id
    plt.close('all')
    print(completed)
    '''
    except TypeError:
        incompleted = "Chưa nhập file chạy " + obj.lab_id
    '''
    return completed, incompleted

def process_HCV_runfile(green_file, yellow_file):
    completed = []
    not_complete = []
    
    try:
        file_1 = xml.dom.minidom.parse(green_file)
        root_1 = file_1.documentElement
        data_1 = {}
        series = root_1.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Hh][Cc][vV]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_1[name] = array

        file_2 = xml.dom.minidom.parse(yellow_file)
        root_2 = file_2.documentElement
        data_2 = {}
        series = root_2.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Hh][Cc][vV]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_2[name] = array
    except AttributeError:
        not_complete.append('File chạy không chứa mẫu tên HCV')

    else:
        for k, name in list(enumerate(data_1)):
            try:
                pt = HCVSample.objects.filter(lab_id=name).order_by('added').first()
                js_data = {"green": data_1[name], "yellow": data_2[name]}
                js = json.dumps(js_data)
                pt.curves = js
                pt.save()

                completed.append(name)
            except(HCVSample.DoesNotExist, ObjectDoesNotExist):

                not_complete.append(name)
            except KeyError:
                not_complete.append('Hai files chạy không khớp nhau')
    for id_ in completed:
        obj = HCVSample.objects.get(lab_id = id_)
        print(obj.pk)
        c,i = create_HCV_image(obj.pk)
    
    return completed, not_complete

def create_HCV_report(pk = None):
    obj = HCVSample.objects.get(pk = pk)
    completed = ""
    incompleted = ""
    try:
        desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
        today_folder =  "D:/Result/" + datetime.now().strftime("%m%d")
        if not os.path.exists(today_folder):
            os.makedirs(today_folder, exist_ok = True)
        try:
            print(obj.sid)
            sid_short = re.search(r"\d+-(\d{4})", obj.sid).groups()[0]

        except (AttributeError, TypeError):
            sid_short = ""
        file_name = today_folder +"/" + sid_short + '-' + no_accent_vietnamese(obj.name.replace(" ",'')) + '-' +obj.lab_id+'.docx'
        pic_file = settings.STATIC_DIR + "/run_img/" + obj.lab_id + '_amp.png'
        if obj.sid == None:
            sid = ""
        else:
            sid = str(obj.sid)
        
        if obj.age != None:
            age = str(obj.age)
        else:
            age = ""
        #start merging
        f_1 = MailMerge(settings.STATIC_DIR + '/word_template/HCV FORM.docx')
        f_1.merge(
            sid = sid, 
            lab_id = str(obj.lab_id), 
            sex = obj.sex, 
            doctor = obj.doctor, 
            addres = obj.address, 
            age = age, 
            name = obj.name, 
            date_receive = obj.added.strftime('%d/%m/%Y'), 
            date_finished = obj.modified.strftime('%d/%m/%Y-%H:%M'), 
            d= datetime.now().strftime("%d"), 
            m= datetime.now().strftime('%m'), 
            y= datetime.now().strftime('%Y'))      
        if obj.result == "DƯỚI NGƯỠNG PHÁT HIỆN":
            f_1.merge(iu = "DƯỚI NGƯỠNG PHÁT HIỆN")
        elif obj.copies != None:        
            try:            
                f_1.merge(
                copies =  "{:.1E}".format(Decimal(obj.copies*obj.std_curve.factor)),
                iu = "{:.1E}".format(Decimal(obj.copies*obj.std_curve.factor/5)),
                log = "{:.1f}".format(math.log10(obj.copies*obj.std_curve.factor)))
            except:
                print('ERROR')

        f_1.write(file_name)

        doc = Document(file_name)
        for paragraph in doc.paragraphs:
            if r"None" in paragraph.text:
                print('YES NONE IS HERE')
                paragraph.text = paragraph.text.strip().replace("None", "")
            if "[image]" in paragraph.text:
                paragraph.text = paragraph.text.strip().replace("[image]", "")

                run = paragraph.add_run()
                run.add_picture(pic_file, width=Cm(15))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(file_name)
        completed = obj.lab_id + " vào " + today_folder
    except FileNotFoundError:
        incompleted = "Chưa có file hình " + obj.lab_id
    except:
        incompleted = obj.lab_id
    return completed, incompleted

#CTNG FUNCTIONS
def create_CTNG_image(pk = None):
    f = open('result/ctng_neg.json')
    ctng_neg = json.load(f)
    cycles = np.arange(1, 41)
    obj = get_object_or_404(CTNGSample, pk = pk)
    curves = json.loads(obj.curves)
    completed = ""
    incompleted = ""
    fig, (ax1,ax2) = plt.subplots(1,2)
    fig.set_size_inches(10, 4)
    
    ax1.plot(cycles, pd.DataFrame(ctng_neg['green']),label = 'N. gonorrhoeae', c = 'blue', marker = "s", markersize = 3)
    ax1.plot(cycles, pd.DataFrame(ctng_neg['yellow']),label = 'Chứng nội', c = 'green', marker = "o", markersize = 3)
    ax1.plot(cycles, pd.DataFrame(ctng_neg['orange']),label = 'C.trachomatis', c = 'red', marker = "<",  markersize = 3)
   
    #ax1.plot(cycles, pd.DataFrame(data_2[item[1]]), label = item[1], c = 'red', linewidth = 2.3)
    #ax1.plot(cycles, pd.DataFrame(data_3[item[1]]), label = 'Chứng nội ' + str(item[1]), c = 'green', alpha = 0.7, marker = '+', markersize = 3)
    ax1.legend()
    ax1.set_xlabel('Cycles')
    ax1.set_ylabel('Normalized Fluorescence')
    ax1.set_ylim(ymax = 0.8)

    ax1.set_title('Mẫu âm tính')
    ax1.grid(linestyle = ':', alpha = 0.5)

    
    ax2.plot(cycles, pd.DataFrame(curves['green']),label = 'N. gonorrhoeae', c = 'blue', marker = "s", markersize = 3)
    ax2.plot(cycles, pd.DataFrame(curves['yellow']),label = 'Chứng nội', c = 'green', marker = "o", markersize = 3)
    ax2.plot(cycles, pd.DataFrame(curves['orange']),label = 'C.trachomatis', c = 'red', marker = "<",  markersize = 3)
    
    ax2.legend()
    ax2.set_title(r"Mẫu " + obj.lab_id)
    ax2.set_xlabel('Cycles')
    ax2.set_ylabel('Normalized Fluorescence')
    ax2.grid(linestyle = ':', alpha = 0.5)
    if ax2.get_ylim()[1] < 0.5:
        ax2.set_ylim(ymax = 0.8)
    fig.tight_layout(h_pad = 3)
    props = dict(boxstyle='round', alpha=0)
    name = "static/run_img/" + obj.lab_id + '-VA.png'
    fig.savefig(name, type = 'png', bbox_inches = 'tight', dpi= 100)
    completed += obj.lab_id
    return completed, incompleted

def process_CTNG_runfile(green_file, orange_file, yellow_file):
    completed = []
    not_complete = []    
    try:
        file_1 = xml.dom.minidom.parse(green_file)
        root_1 = file_1.documentElement
        data_1 = {}
        series = root_1.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Cc][Tt][Nn][Gg]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_1[name] = array

        file_2 = xml.dom.minidom.parse(yellow_file)
        root_2 = file_2.documentElement
        data_2 = {}
        series = root_2.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Cc][Tt][Nn][Gg]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_2[name] = array
        
        file_3 = xml.dom.minidom.parse(orange_file)
        root_3 = file_3.documentElement
        data_3 = {}
        series = root_3.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Cc][Tt][Nn][Gg]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_3[name] = array
    except AttributeError:
        not_complete.append('File chạy không chứa mẫu tên CTNG')

    else:
        for k, name in list(enumerate(data_1)):
            try:
                pt = CTNGSample.objects.filter(lab_id=name).order_by('added').first()
                js_data = {"green": data_1[name], "yellow": data_2[name], "orange": data_3[name]}
                js = json.dumps(js_data)
                pt.curves = js
                pt.save()

                completed.append(name)
            except(CTNGSample.DoesNotExist, ObjectDoesNotExist):

                not_complete.append(name)
            except KeyError:
                not_complete.append('Hai files chạy không khớp nhau')
    for id_ in completed:
        obj = CTNGSample.objects.get(lab_id = id_)
        print(obj.pk)
        c,i = create_CTNG_image(obj.pk)
    
    return completed, not_complete

def create_CTNG_report(pk = None):
    obj = CTNGSample.objects.get(pk = pk)
    completed = ""
    incompleted = ""
    #try:
    desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
    today_folder =  "D:/Result/" + datetime.now().strftime("%m%d")
    if not os.path.exists(today_folder):
        os.makedirs(today_folder, exist_ok = True)
    try:
        print(obj.sid)
        sid_short = re.search(r"\d+-(\d{4})", obj.sid).groups()[0]

    except (AttributeError, TypeError):
        sid_short = ""
    file_name = today_folder +"/" + sid_short + '-' + no_accent_vietnamese(obj.name.replace(" ",'')) + '-' +obj.lab_id+'.docx'
    pic_file = settings.STATIC_DIR + "/run_img/" + obj.lab_id + '-VA.png'
    if obj.sid == None:
        sid = ""
    else:
        sid = str(obj.sid)
    
    if obj.age != None:
        age = str(obj.age)
    else:
        age = ""
    result_ct_neg = ''
    result_ct_pos = ''
    result_ng_neg = ''
    result_ng_pos = ''
    if obj.result_ct == 'ÂM TÍNH':
        result_ct_neg = 'ÂM TÍNH'
    else:
        result_ct_pos = 'DƯƠNG TÍNH'

    if obj.result_ng == 'ÂM TÍNH':
        result_ng_neg = 'ÂM TÍNH'
    else:
        result_ng_pos = 'DƯƠNG TÍNH'
    print(obj.modified.strftime('%d/%m/%Y-%H:%M'))
    #start merging
    f_1 = MailMerge(settings.STATIC_DIR + '/word_template/CTNG FORM.docx')
    f_1.merge(
        sid = sid, 
        lab_id = str(obj.lab_id), 
        sex = obj.sex, 
        doctor = obj.doctor, 
        addres = obj.address, 
        age = age, 
        name = obj.name, 
        date_receive = obj.added.strftime('%d/%m/%Y'), 
        date_finished = obj.modified.strftime('%d/%m/%Y-%H:%M'), 
        d= datetime.now().strftime("%d"), 
        m= datetime.now().strftime('%m'), 
        y= datetime.now().strftime('%Y'),
        result_ct_neg = result_ct_neg,
        result_ct_pos = result_ct_pos,
        result_ng_pos = result_ng_pos,
        result_ng_neg = result_ng_neg,
        sample_type = obj.sample_type
        )     


    f_1.write(file_name)

    doc = Document(file_name)
    for paragraph in doc.paragraphs:
        if r"None" in paragraph.text:
            print('YES NONE IS HERE')
            paragraph.text = paragraph.text.strip().replace("None", "")
        if "[image]" in paragraph.text:
            paragraph.text = paragraph.text.strip().replace("[image]", "")

            run = paragraph.add_run()
            run.add_picture(pic_file, width=Cm(15))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(file_name)
    completed = obj.lab_id + " vào " + today_folder
    '''
    except FileNotFoundError:
        incompleted = "Chưa có file hình " + obj.lab_id
    except:
        incompleted = obj.lab_id'''
    return completed, incompleted

#HPV FUNCTIONS
def create_HPV_image(pk = None):
    
    cycles = np.arange(1, 41)
    obj = get_object_or_404(HPVSample, pk = pk)
    curves = json.loads(obj.curves)
    completed = ""
    incompleted = ""
    if obj.test_kit == "KT":
        f = open('result/hpv_neg_kt.json')
        hpv_neg = json.load(f)
        fig, (ax1,ax2) = plt.subplots(1,2)
        fig.set_size_inches(10, 4)
        
        ax1.plot(cycles, pd.DataFrame(hpv_neg['green']),label = 'HPV-16', c = 'green', marker = "s", markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['yellow']),label = 'HPV-18', c = '#2142FF', marker = "o", markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['orange']),label = 'HPV High Risk', c = '#FF8600', marker = "<",  markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['red']),label = 'Chứng nội', c = '#E80D0C', marker = 'x', markersize = 3)

        ax1.legend()
        ax1.set_xlabel('Cycles')
        ax1.set_ylabel('Normalized Fluorescence')
        ax1.set_ylim(ymax = 1.0)

        ax1.set_title('Mẫu âm tính')
        ax1.grid(linestyle = ':', alpha = 0.5)
        
        ax2.plot(cycles, pd.DataFrame(curves['green']),label = 'HPV-16', c = 'green', marker = "s", markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['yellow']),label = 'HPV-18', c = '#2142FF', marker = "o", markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['orange']),label = 'HPV High Risk', c = '#FF8600', marker = "<",  markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['red']),label = 'Chứng nội', c = '#E80D0C', marker = 'x', markersize = 3)
        ax2.legend()
        ax2.set_title(r"Mẫu " + obj.lab_id)
        ax2.set_xlabel('Cycles')
        ax2.set_ylabel('Normalized Fluorescence')
        ax2.grid(linestyle = ':', alpha = 0.5)
        
        if ax2.get_ylim()[1] < 0.5:
            ax2.set_ylim(ymax = 0.5)
        
        fig.tight_layout(h_pad = 3)
        props = dict(boxstyle='round', alpha=0)
        name = "static/run_img/" + obj.lab_id+ '-KT.png'
        fig.savefig(name, type = 'png', bbox_inches = 'tight', dpi= 100)
        completed += obj.lab_id
        f.close()

    elif obj.test_kit == "VA":
        fig = plt.figure()    
        plt.plot(cycles, pd.DataFrame(curves['green']),label = 'HPV-DNA', c = 'red', marker = "s", markersize = 3)
        plt.plot(cycles, pd.DataFrame(curves['yellow']),label = 'Chứng nội', c = 'green', marker = "o", markersize = 3)
        
        plt.legend()
        plt.title(r"Mẫu " + obj.lab_id)
        plt.xlabel('Cycles')
        plt.ylabel('Normalized Fluorescence')
        plt.grid(linestyle = ':', alpha = 0.5)
        #ax2.set_ylim()
        fig.tight_layout(h_pad = 3)
        props = dict(boxstyle='round', alpha=0)
        name = "static/run_img/" + obj.lab_id + '-VA.png'
        fig.savefig(name, type = 'png', bbox_inches = 'tight', dpi= 100)
        completed += obj.lab_id
    elif obj.test_kit == "KT_611":
        f = open('result/hpv_neg_kt.json')
        hpv_neg = json.load(f)
        fig, (ax1,ax2) = plt.subplots(1,2)
        fig.set_size_inches(10, 4)
        
        ax1.plot(cycles, pd.DataFrame(hpv_neg['green']),label = 'HPV-16', c = 'green', marker = "s", markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['yellow']),label = 'HPV-18', c = '#2142FF', marker = "o", markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['orange']),label = 'HPV High Risk', c = '#FF8600', marker = "<",  markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['red']),label = 'Chứng nội', c = '#E80D0C', marker = 'x', markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['green_2']),label = 'HPV-6', c = '#460656', marker = "4",  markersize = 3)
        ax1.plot(cycles, pd.DataFrame(hpv_neg['yellow_2']),label = 'HPV-11', c = '#10073a', marker = "d",  markersize = 3)

        ax1.legend()
        ax1.set_xlabel('Cycles')
        ax1.set_ylabel('Normalized Fluorescence')
        ax1.set_ylim(ymax = 1.0)

        ax1.set_title('Mẫu âm tính')
        ax1.grid(linestyle = ':', alpha = 0.5)
        
        ax2.plot(cycles, pd.DataFrame(curves['green']),label = 'HPV-16', c = 'green', marker = "s", markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['yellow']),label = 'HPV-18', c = '#2142FF', marker = "o", markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['orange']),label = 'HPV High Risk', c = '#FF8600', marker = "<",  markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['red']),label = 'Chứng nội', c = '#E80D0C', marker = 'x', markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['green_2']),label = 'HPV-6', c = '#460656', marker = "4",  markersize = 3)
        ax2.plot(cycles, pd.DataFrame(curves['yellow_2']),label = 'HPV-11', c = '#10073a', marker = "d",  markersize = 3)
        ax2.legend()
        ax2.set_title(r"Mẫu " + obj.lab_id)
        ax2.set_xlabel('Cycles')
        ax2.set_ylabel('Normalized Fluorescence')
        ax2.grid(linestyle = ':', alpha = 0.5)
        
        if ax2.get_ylim()[1] < 0.5:
            ax2.set_ylim(ymax = 0.5)
        
        fig.tight_layout(h_pad = 3)
        props = dict(boxstyle='round', alpha=0)
        name = "static/run_img/" + obj.lab_id+ '-KTPLUS.png'
        fig.savefig(name, type = 'png', bbox_inches = 'tight', dpi= 100)
        completed += obj.lab_id
        f.close()

    return completed, incompleted

def process_HPV_runfile(green_file, yellow_file, orange_file = None, red_file = None, green_file_2 = None, yellow_file_2 = None):
    completed = []
    not_complete = []    
    try:
        file_1 = xml.dom.minidom.parse(green_file)
        root_1 = file_1.documentElement
        data_1 = {}
        series = root_1.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Hh][Pp][Vv]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_1[name] = array

        file_2 = xml.dom.minidom.parse(yellow_file)
        root_2 = file_2.documentElement
        data_2 = {}
        series = root_2.getElementsByTagName('series')
        for serie in series:
            name = re.search(r'([Hh][Pp][Vv]-\S*)',
                             serie.getAttribute('title')).group(0)
            array = []
            for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                array.append(float(point.getAttribute('Y')))
            data_2[name] = array
        if orange_file != None:
            file_3 = xml.dom.minidom.parse(orange_file)
            root_3 = file_3.documentElement
            data_3 = {}
            series = root_3.getElementsByTagName('series')
            for serie in series:
                name = re.search(r'([Hh][Pp][Vv]-\S*)',
                                serie.getAttribute('title')).group(0)
                array = []
                for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                    array.append(float(point.getAttribute('Y')))
                data_3[name] = array
        if red_file != None:
            file_4 = xml.dom.minidom.parse(red_file)
            root_4 = file_4.documentElement
            data_4 = {}
            series = root_4.getElementsByTagName('series')
            for serie in series:
                name = re.search(r'([Hh][Pp][Vv]-\S*)',
                                serie.getAttribute('title')).group(0)
                array = []
                for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                    array.append(float(point.getAttribute('Y')))
                data_4[name] = array
        if green_file_2 != None:
            file_5 = xml.dom.minidom.parse(green_file_2)
            root_5 = file_5.documentElement
            data_5 = {}
            series = root_5.getElementsByTagName('series')
            for serie in series:
                name = re.search(r'([Hh][Pp][Vv]-\S*)',
                                serie.getAttribute('title')).group(0)
                array = []
                for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                    array.append(float(point.getAttribute('Y')))
                data_5[name] = array
        
        if yellow_file_2 != None:
            file_6 = xml.dom.minidom.parse(yellow_file_2)
            root_6 = file_6.documentElement
            data_6 = {}
            series = root_6.getElementsByTagName('series')
            for serie in series:
                name = re.search(r'([Hh][Pp][Vv]-\S*)',
                                serie.getAttribute('title')).group(0)
                array = []
                for point in serie.getElementsByTagName('points')[0].getElementsByTagName('point'):
                    array.append(float(point.getAttribute('Y')))
                data_6[name] = array
    except AttributeError:
        not_complete.append('File chạy không chứa mẫu tên HPV')

    else:
        for k, name in list(enumerate(data_1)):
            try:
                pt = HPVSample.objects.filter(lab_id=name).order_by('added').first()
                js_data = {"green": data_1[name], "yellow": data_2[name]}
                if orange_file != None:
                    js_data['orange'] = data_3[name]
                if red_file != None:
                    js_data['red'] = data_4[name]
                if green_file_2 != None:
                    js_data['green_2'] = data_5[name]
                if yellow_file_2 != None:
                    js_data['yellow_2'] = data_6[name]
                js = json.dumps(js_data)

                pt.curves = js

                pt.save()

                completed.append(name)
            except(HPVSample.DoesNotExist, ObjectDoesNotExist):

                not_complete.append(name)
            except KeyError:
                not_complete.append('Hai files chạy không khớp nhau')
    for id_ in completed:
        obj = HPVSample.objects.get(lab_id = id_)
        print(obj.pk)
        c,i = create_HPV_image(obj.pk)
    
    return completed, not_complete

def create_HPV_report(pk = None):
    obj = HPVSample.objects.get(pk = pk)
    completed = ""
    incompleted = ""
    #try:
    desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
    today_folder =  "D:/Result/" + datetime.now().strftime("%m%d")
    if not os.path.exists(today_folder):
        os.makedirs(today_folder, exist_ok = True)
    try:
        print(obj.sid)
        sid_short = re.search(r"\d+-(\d{4})", obj.sid).groups()[0]

    except (AttributeError, TypeError):
        sid_short = ""
    file_name = today_folder +"/" + sid_short + '-' + no_accent_vietnamese(obj.name.replace(" ",'')) + '-' +obj.lab_id+'.docx'
    
    if obj.sid == None:
        sid = ""
    else:
        sid = str(obj.sid)
    if obj.age != None:
        age = str(obj.age)
    else:
        age = ""
    if obj.test_kit == "KT":
        pic_file = "static/run_img/" + obj.lab_id + "-KT.png"
        result_16_neg = ""
        result_16_pos = ""
        result_18_neg = ""
        result_18_pos = ""
        result_hr_neg = ""
        result_hr_pos = ""

        if obj.result_16_kt == "DƯƠNG TÍNH":
            result_16_pos = "DƯƠNG TÍNH"
        else: 
            result_16_neg = "ÂM TÍNH"

        if obj.result_18_kt == "DƯƠNG TÍNH":
            result_18_pos = "DƯƠNG TÍNH"
        else: 
            result_18_neg = "ÂM TÍNH"

        if obj.result_hr_kt == "DƯƠNG TÍNH":
            result_hr_pos = "DƯƠNG TÍNH"
        else: 
            result_hr_neg = "ÂM TÍNH"
        f_1 = MailMerge('static/word_template/HPV FORM.docx')
        f_1.merge(
            sid = obj.sid,
            lab_id = " " + obj.lab_id, 
            result_18_neg = result_18_neg,
            result_18_pos = result_18_pos,
            result_16_pos = result_16_pos,
            result_16_neg = result_16_neg,
            result_hr_pos = result_hr_pos,
            result_hr_neg = result_hr_neg,
            sex = obj.sex,
            doctor =obj.doctor, 
            address = obj.address, 
            age = age, 
            name = obj.name, 
            date_receive = obj.added.strftime('%d/%m/%Y'), 
            date_finished = obj.modified.strftime('%d/%m/%Y-%H:%M'), 
            d= datetime.now().strftime("%d"), 
            m= datetime.now().strftime('%m'), 
            y= datetime.now().strftime('%Y'))
    
        f_1.write(file_name)
        
        

        doc = Document(file_name)
        for paragraph in doc.paragraphs:
            if "[image]" in paragraph.text:
                paragraph.text = paragraph.text.strip().replace("[image]", "")

                run = paragraph.add_run()
                run.add_picture(pic_file, width=Cm(15))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.save(file_name)
        completed = obj.lab_id + " vào " + today_folder

    elif obj.test_kit == "VA":
        pic_file = "static/run_img/" + obj.lab_id + "-VA.png"
        result_hpv_neg = ""
        result_hpv_pos = ""
        result_geno_pos = ""
        result_geno_neg = ""

        if obj.result_qual_va == "DƯƠNG TÍNH":
            result_hpv_pos = "DƯƠNG TÍNH"
            result_geno_pos = obj.result_type_va
        else:
            result_hpv_neg = "ÂM TÍNH"        
            result_geno_neg = "ÂM TÍNH"
        f_1 = MailMerge('static/word_template/HPV FORM VA.docx')
        f_1.merge(
        sid = obj.sid,
        lab_id = " " + obj.lab_id, 
        result_hpv_neg = result_hpv_neg,
        result_hpv_pos = result_hpv_pos,
        result_geno_pos = result_geno_pos,
        result_geno_neg = result_geno_neg,
        sample_type = obj.sample_type,
        sex = obj.sex, 
        doctor = obj.doctor, 
        address = obj.address, 
        age = age, 
        name = obj.name, 
        date_receive = obj.added.strftime('%d/%m/%Y'), 
        date_finished = obj.modified.strftime('%d/%m/%Y-%H:%M'), 
        d= datetime.now().strftime("%d"), 
        m= datetime.now().strftime('%m'), 
        y= datetime.now().strftime('%Y'))
        f_1.write(file_name)      
        
        doc = Document(file_name)
        for paragraph in doc.paragraphs:
            if "[image]" in paragraph.text:
                paragraph.text = paragraph.text.strip().replace("[image]", "")
                run = paragraph.add_run()
                run.add_picture(pic_file, width=Cm(8.5))
                run.add_picture('static/run_img/valayout.gif', width=Cm(3.8))
                if obj.result_qual_va == "ÂM TÍNH":                    
                    run.add_picture('static/run_img/hpv neg.jpg', width = Cm(3.8))
                elif obj.result_qual_va == "DƯƠNG TÍNH":
                    run.add_picture(obj.img_va, width = Cm(3.8))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.save(file_name)
        completed = obj.lab_id + " vào " + today_folder

    if obj.test_kit == "KT_611":
        pic_file = "static/run_img/" + obj.lab_id + "-KTPLUS.png"
        result_16_neg = ""
        result_16_pos = ""
        result_18_neg = ""
        result_18_pos = ""
        result_hr_neg = ""
        result_hr_pos = ""
        result_6_neg = ""
        result_6_pos = ""
        result_11_neg = ""
        result_11_pos = ""

        if obj.result_16_kt == "DƯƠNG TÍNH":
            result_16_pos = "DƯƠNG TÍNH"
        else: 
            result_16_neg = "ÂM TÍNH"

        if obj.result_18_kt == "DƯƠNG TÍNH":
            result_18_pos = "DƯƠNG TÍNH"
        else: 
            result_18_neg = "ÂM TÍNH"

        if obj.result_hr_kt == "DƯƠNG TÍNH":
            result_hr_pos = "DƯƠNG TÍNH"
        else: 
            result_hr_neg = "ÂM TÍNH"

        if obj.result_6_kt == "DƯƠNG TÍNH":
            result_6_pos = "DƯƠNG TÍNH"
        else: 
            result_6_neg = "ÂM TÍNH"

        if obj.result_11_kt == "DƯƠNG TÍNH":
            result_11_pos = "DƯƠNG TÍNH"
        else: 
            result_11_neg = "ÂM TÍNH" 
        f_1 = MailMerge('static/word_template/HPV611 FORM.docx')
        f_1.merge(
            sid = obj.sid,
            lab_id = " " + obj.lab_id, 
            result_18_neg = result_18_neg,
            result_18_pos = result_18_pos,
            result_16_pos = result_16_pos,
            result_16_neg = result_16_neg,
            result_hr_pos = result_hr_pos,
            result_hr_neg = result_hr_neg,
            result_6_pos = result_6_pos,
            result_6_neg = result_6_neg,
            result_11_pos = result_11_pos,
            result_11_neg = result_11_neg,
            sex = obj.sex,
            doctor =obj.doctor, 
            address = obj.address, 
            age = age, 
            name = obj.name, 
            date_receive = obj.added.strftime('%d/%m/%Y'), 
            date_finished = obj.modified.strftime('%d/%m/%Y-%H:%M'), 
            d= datetime.now().strftime("%d"), 
            m= datetime.now().strftime('%m'), 
            y= datetime.now().strftime('%Y'))
    
        f_1.write(file_name)
        
        

        doc = Document(file_name)
        for paragraph in doc.paragraphs:
            if "[image]" in paragraph.text:
                paragraph.text = paragraph.text.strip().replace("[image]", "")

                run = paragraph.add_run()
                run.add_picture(pic_file, width=Cm(15))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.save(file_name)
        completed = obj.lab_id + " vào " + today_folder
        
    '''
    except FileNotFoundError:
        incompleted = "Chưa có file hình " + obj.lab_id
    except:
        incompleted = obj.lab_id'''
    return completed, incompleted

def create_red_run():
	arr = [-0.0307022044209794, -0.0280913091964497, -0.0254804139719199, -0.0228695187473902, -0.0204225687144872, -0.0180861506800482, -0.0157451881149236, -0.0132951264573411, -0.010757233874206, -0.00831366353130584, -0.00600839059388634, -0.00393327891278045, -0.00194970935807219, 0.000107940355260816, 0.00241048122940581, 0.00488721428384435, 0.00740852208138584, 0.00981269750245809, 0.0119529693744361, 0.0138114563660028, 0.0155742208295958, 0.0177146635781359, 0.0208902330238241, 0.0262995110617809, 0.0359125335872005, 0.0530687239571791, 0.0822455264825406, 0.128130786654408, 0.193579836149974, 0.278116153074254, 0.376969566008115, 0.481960430558091, 0.583440025861203, 0.673036977339407, 0.745318703643666, 0.798278306493034, 0.835157142336628, 0.860241914617969, 0.885326686899309, 0.910411459180649]
	arr_2 = [num + random.uniform(-0.005, 0.005) for num in arr]
	return arr_2

def populate_HPV_result():
    data = {
        'green': np.random.uniform(-0.005, 0.005, 40).tolist(),
        'yellow': np.random.uniform(0.01, 0.015, 40).tolist(),
        'orange': np.random.uniform(-0.005, 0.005, 40).tolist(),
        'red': create_red_run(),
        'green_2': np.random.uniform(0.005, 0.009, 40).tolist(),
        'yellow_2': np.random.uniform(-0.005, 0.005, 40).tolist(),
    }
    data = json.dumps(data)
    return data

#HCVGENO FUNCTIONS
def create_HCV_geno_img(pk = None):
    cycles = np.arange(1, 41)
    obj = get_object_or_404(HCVGenoSample, pk = pk)
    curves = json.loads(obj.curves)
    print(curves)
    fig = plt.figure(figsize=(6,4))
    for k, v in curves.items():
        plt.plot(pd.DataFrame(cycles), pd.DataFrame(v), label = k)
    plt.legend()
    plt.xlabel('Cycles')
    plt.ylabel('Normalized Fluorescence')
    plt.title('HCV Genotype ' + obj.sample.lab_id)
    plt.tight_layout()
    plt.grid(linestyle = ':', alpha = 0.5)
    plt.ylim(ymax = 0.5)
    name = "static/run_img/" + obj.sample.lab_id+ '_genotype.png'
    fig.savefig(name, type = 'png',  bbox_inches = 'tight', dpi= 100)

def populate_HCV_geno(pk = None, genotype = None):
    obj = get_object_or_404(HCVGenoSample, pk = pk)
    pos_curve = [
        [0.00850570344926657, 0.0078039564298498, 0.00710220941043302, 0.00640046239101624, 0.00568994087860778, 0.00489241257277972, 0.00399357651197132, 0.00292907272596679, 0.00161023608669879, 6.84841742102838e-05, -0.00148545435602104, -0.00284604624880491, -0.00387280882329724, -0.00440839513446598, -0.0044948203670799, -0.00428503435847069, -0.00410809596902337, -0.00407186043452459, -0.00425302912335318, -0.00449826365855962, -0.00460999358177076, -0.00402071597194069, -0.00201399096000546, 0.00258844148264283, 0.0116024493262048, 0.0281244084594061, 0.0563529369845273, 0.100685476638644, 0.16391388821728, 0.245745610249846, 0.342239078633314, 0.446459121929933, 0.550262198191539, 0.646441767340846, 0.730430631212083, 0.8004382442779, 0.85798310762489, 0.905477814406338, 0.952972521187785, 1.00046722796923], 
        [0.00502462450768613, 0.00487385317619148, 0.00472308184469683, 0.00457231051320218, 0.00428039626220306, 0.00378689643619647, 0.00297207118994891, 0.00186533988061461, 0.000524466856939555, -0.000896970743876267, -0.00222939295154554, -0.00323111911803258, -0.00374705318461728, -0.0035769837402781, -0.00245617279694779, 0.000358937447574895, 0.00635862874946144, 0.0182447019018473, 0.0400006242907231, 0.0765452210997706, 0.132180102605265, 0.208660908770037, 0.304036260539178, 0.412746350571052, 0.526802392347728, 0.637714893585508, 0.738675611951806, 0.825600900199354, 0.897014951939224, 0.95324195071379, 0.995685218938614, 1.0262866271825, 1.0471556627893, 1.0603642891248, 1.06781067539479, 1.07120545244924, 1.07204880575064, 1.07145557276413, 1.07086233977761, 1.0702691067911],
        [0.0161795612560942, 0.0120155841329027, 0.00785160700971121, 0.00368762988651969, 0.000213003456402615, -0.00247742207649026, -0.00421461302634069, -0.00493388135827954, -0.00423195606175224, -0.00113066459027797, 0.00630902979206391, 0.0212599767760123, 0.0482165683116318, 0.0919709740298346, 0.155909152621082, 0.240239797510015, 0.341416715982471, 0.452603268440953, 0.565370730047073, 0.671849098832084, 0.766500950167862, 0.846552631487183, 0.911417101653138, 0.961918912730762, 0.999619042161527, 1.02641459973751, 1.04425342492389, 1.05502271223437, 1.06049504635863, 1.06228465625955, 1.06174704249937, 1.0599116185633, 1.05747503022703, 1.05483899343191, 1.05217728670518, 1.04953149413247, 1.04690172377949, 1.0442877827471, 1.04167384171471, 1.03905990068232],
        [0.00764131353300749, 0.00654714424108294, 0.00545297494915838, 0.00435880565723382, 0.00340473585061176, 0.00263293364665634, 0.00209418423093339, 0.0017192439609039, 0.00144374561639838, 0.0010933048910998, 0.000626421573827717, -1.21044252290434e-05, -0.000713663125724892, -0.00145615169373934, -0.00216674874565307, -0.00285736644049172, -0.0035134438857042, -0.00412797615343818, -0.00458665788570905, -0.00469188434281159, -0.00415013214919175, -0.00243671414689326, 0.00144059784842488, 0.00925094136938969, 0.0239101257215007, 0.0495430349411185, 0.0907307191784924, 0.150632653756885, 0.229473883015106, 0.323844393112966, 0.427324090540779, 0.531799606764026, 0.629772773164867, 0.715960934319227, 0.787967526911495, 0.845597560139515, 0.891396514775706, 0.928264864902606, 0.965133215029506, 1.00200156515641],
        [0.00155159386365703, 0.00197749372923349, 0.00240339359480994, 0.00282929346038639, 0.0033325979601831, 0.00380273443061356, 0.00418510400172643, 0.00425115978446299, 0.00394498967459716, 0.00320367614425436, 0.00216490053026159, 0.000916200558005723, -0.000259807203427046, -0.00122494386733158, -0.00184298103388501, -0.00216461711335479, -0.00223363119300147, -0.0022142763477148, -0.00217440579750342, -0.00211178736184093, -0.0019227746473467, -0.00156146517479844, -0.00102682361493124, -0.000380509000412519, 0.00031626142650768, 0.00104476633619514, 0.00193385627208348, 0.00333803606982546, 0.00596767886820119, 0.0110336590034834, 0.0205194427542181, 0.037332523711711, 0.0650395179459554, 0.106771675734603, 0.163726509898416, 0.234848972102986, 0.31364143547254, 0.393824849440726, 0.474008263408912, 0.554191677377098],
        [-0.0169705795949503, -0.0156024152000589, -0.0142342508051674, -0.012866086410276, -0.0114367931551777, -0.0100459451299411, -0.00879462418383612, -0.00785755644981566, -0.00726012638303516, -0.00690997374057838, -0.00659979646748731, -0.00611782494491142, -0.00532182032408628, -0.00418580178020347, -0.00278927454321915, -0.00129587161422289, 0.000149023083325958, 0.00148476540623629, 0.00281418658239232, 0.00435221577249169, 0.00653224155469306, 0.0100171235052257, 0.016012392646167, 0.0266081520299443, 0.0455548099517164, 0.0784938086558523, 0.132061103093003, 0.211175278169349, 0.316426715664943, 0.44274275246253, 0.580103761352103, 0.715856371695345, 0.838302111171601, 0.939577776833788, 1.01665416754649, 1.07033055114702, 1.10592512836174, 1.12894323618046, 1.15196134399918, 1.1749794518179],
        [-0.0295764397494571, -0.0265428582631206, -0.0235092767767841, -0.0204756952904477, -0.0177453288111408, -0.0153637164243583, -0.0133579007492934, -0.0116272167733756, -0.0101190874937525, -0.00875782062190659, -0.007429925887978, -0.00599857499799743, -0.00443436482004375, -0.00269303939853929, -0.000910694178766116, 0.000840270378289411, 0.00242194638379194, 0.00397444700489578, 0.00557724263017649, 0.00753136666111084, 0.0101675093968287, 0.0143180231056752, 0.0214397518662162, 0.0343743612006338, 0.0577148648611245, 0.0977818489317352, 0.160833449502323, 0.250262819575461, 0.364414424238157, 0.496176334179355, 0.634339046896748, 0.766371040020982, 0.88191516476847, 0.974953314246265, 1.04403090569304, 1.09099819001975, 1.1212858808044, 1.14023787791987, 1.15918987503534, 1.17814187215081],
    ]

    neg_curve = [
        [0.00177181534508569, 0.0020155586726796, 0.00225930200027352, 0.00250304532786743, 0.00260180406888613, 0.00253565033414153, 0.00231495100933682, 0.00202160703733897, 0.00174776967058403, 0.00157061577659768, 0.00151280152953814, 0.00155068020791689, 0.00164129660787523, 0.00175784311871588, 0.00189545978241261, 0.00205152524201968, 0.00223904864358243, 0.00241965713674679, 0.00254360080627132, 0.00253739565646651, 0.00242513001017782, 0.00222152721941515, 0.0020049037810318, 0.00181448294497314, 0.00171403754563648, 0.00167856979607729, 0.00169716078692871, 0.00173424693232994, 0.00178104489372263, 0.00181249692396443, 0.00184170287673496, 0.00187236123321268, 0.00188995488510716, 0.00184225163038543, 0.00165720729681364, 0.0011997813048383, 0.000465158160022779, -0.000527311437151558, -0.00151978103432589, -0.00251225063150023],
        [0.000793542544634962, 0.000699194202842018, 0.000604845861049073, 0.000510497519256128, 0.000457046864641867, 0.000497930580176961, 0.000636103183753825, 0.000814812004362063, 0.0010117531044736, 0.00117756945104678, 0.00126270058496392, 0.00128767954749935, 0.00136544481085637, 0.00157162096029459, 0.00188161258014459, 0.00219310497074446, 0.00243168254384872, 0.00251077329685085, 0.00239928100951582, 0.0021354426210283, 0.00182014675944385, 0.00147937575214916, 0.0011007234068629, 0.000714622245645761, 0.000424075557804575, 0.000325166317608381, 0.000438493794340827, 0.000746063591817322, 0.00111490237692965, 0.00136433725140738, 0.00133608758618949, 0.00103290157061571, 0.000550141521491444, 1.1062302607358e-05, -0.000459802390779148, -0.000746820241837, -0.000848733260213468, -0.000812759889197535, -0.000776786518181602, -0.00074081314716567],
        [0.00722452596994101, 0.00658292952762543, 0.00594133308530984, 0.00529973664299425, 0.00461499408096679, 0.00390947805154264, 0.00317358492882926, 0.00246049119701039, 0.0017605493419433, 0.00112308603038783, 0.000521647390420701, 1.97302729704069e-05, -0.000367596611780924, -0.000554755280711511, -0.000571038218738246, -0.000411573970504676, -0.000163010063716499, 0.000114940202101304, 0.000312256236159919, 0.000421525123161911, 0.000447976527456239, 0.000439898348472064, 0.000408820535102497, 0.000389987799649279, 0.000397830320786141, 0.000472719594996172, 0.000670725583090654, 0.00105660422952779, 0.0016682655291382, 0.00250295901599319, 0.00350163046009785, 0.00455621439787068, 0.00557847069856017, 0.00654072465835789, 0.00746501169607506, 0.00839555211146239, 0.00936509128444986, 0.0104000506902711, 0.0114350100960924, 0.0124699695019136],
        [-0.00769373566661931, -0.00687794155446266, -0.006062147442306, -0.00524635333014935, -0.00463215298336626, -0.00397613391260497, -0.0030679076615831, -0.00153239251375544, 0.000542242116147399, 0.00287520424280402, 0.00492707799012852, 0.00642870700237215, 0.00722812764589711, 0.00767541447743995, 0.00815334900075576, 0.0091397731162178, 0.0107241847892455, 0.0128329330385502, 0.0151530234138094, 0.0174071285464983, 0.0193125850424531, 0.0207993454731779, 0.0219585860992144, 0.0229756287162811, 0.023939079876524, 0.0249479844935027, 0.0260287882618872, 0.0271979485329428, 0.0284129073117976, 0.0296890555230419, 0.0310479406991058, 0.0325206353689247, 0.0341318680960281, 0.0358365701679174, 0.0375808460774254, 0.0392996457729027, 0.0409746197822791, 0.0425512158156093, 0.0441278118489396, 0.0457044078822698],
        [4.42833126597657e-05, 0.000591277555937628, 0.00113827179921549, 0.00168526604249335, 0.00200387323111672, 0.00206373700065954, 0.00190037347236453, 0.00164146449076324, 0.00143551326049078, 0.00136991972560593, 0.00146881523358262, 0.00161732188575579, 0.00175587991955913, 0.00181664217449778, 0.00186133004424939, 0.00187525184297803, 0.00195168380440802, 0.00204012455239406, 0.00212924400126185, 0.00210832122541333, 0.00202710200851823, 0.00190048922426652, 0.00177457921219708, 0.00161378200435335, 0.00142666641301627, 0.00118033462101222, 0.000855332372904494, 0.000464629260559332, 0.00012778323515754, -7.84967443713055e-05, -9.45727704515805e-05, 4.84530078791794e-05, 0.000395536964169042, 0.00077383667525066, 0.00100702358123738, 0.000802730068233677, 0.000185250542765573, -0.000901342841213666, -0.0019879362251929, -0.00307452960917214],
        [-0.000962654114823975, -0.000206659230233515, 0.000549335654356945, 0.0013053305389474, 0.00181462805737796, 0.00208739986951407, 0.00213073841226194, 0.00205831426679872, 0.00200164231455627, 0.00207235166652158, 0.00226621381419514, 0.00253304070889635, 0.00276958014383728, 0.00288753754343655, 0.00278629336426269, 0.00250633744981155, 0.00211002198941478, 0.00176434903840493, 0.00149615898468187, 0.00134746177297123, 0.00119566361195735, 0.00100917671349377, 0.000710108116444127, 0.000358729182333695, 2.58006494790241e-05, -0.000180354690300786, -0.000186837529276773, 7.83015221591946e-07, 0.000356190564280117, 0.00073507154405946, 0.00104366151773338, 0.00113491309943366, 0.000928266677557289, 0.000270086031820232, -0.000913780763054217, -0.00277145143455675, -0.00521477078548308, -0.00816266502290132, -0.0111105592603196, -0.0140584534977378]
    ]
    dt = {}
    if genotype == "TYPE1" or genotype == "TYPE 1":
        dt = {
            "TYPE 1": pos_curve[random.randint(0, len(pos_curve))],
            "TYPE 2": neg_curve[random.randint(0, len(pos_curve))],
            "TYPE 6": neg_curve[random.randint(0, len(pos_curve))]
        }
    elif genotype == "TYPE2" or genotype == "TYPE 2":
        dt = {
            "TYPE 2": pos_curve[random.randint(0, len(pos_curve))],
            "TYPE 1": neg_curve[random.randint(0, len(pos_curve))],
            "TYPE 6": neg_curve[random.randint(0, len(pos_curve))]
        }
    elif genotype == "TYPE6" or genotype == "TYPE 6":
        dt = {
            "TYPE 6": pos_curve[random.randint(0, len(pos_curve))],
            "TYPE 2": neg_curve[random.randint(0, len(pos_curve))],
            "TYPE 1": neg_curve[random.randint(0, len(pos_curve))]
        }
    obj.curves = json.dumps(dt)
    obj.save()

def process_HCV_geno_runfile(green_file, pk = None):
    data = {}
    obj = get_object_or_404(HCVGenoSample, pk = pk)
    
    f = xml.dom.minidom.parse(green_file)
    root = f.documentElement
    
    series = root.getElementsByTagName('series')
    for s in series:
        points = s.getElementsByTagName('points')[0]
        array = []
        for point in points.getElementsByTagName('point'):
            name = point.getAttribute('text')
            y = point.getAttribute('Y')
            array.append(float(y))
        data[name] = array

    obj.curves = json.dumps(data)
    obj.save()
    create_HCV_geno_img(pk = pk)
    
def create_HCV_geno_report(pk = None):
    obj_geno = HCVGenoSample.objects.get(pk = pk)
    obj = obj_geno.sample
    completed = ""
    incompleted = ""
    try:
        desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
        today_folder =  "D:/Result/" + datetime.now().strftime("%m%d")
        if not os.path.exists(today_folder):
            os.makedirs(today_folder, exist_ok = True)
        try:
            print(obj.sid)
            sid_short = re.search(r"\d+-(\d{4})", obj.sid).groups()[0]

        except (AttributeError, TypeError):
            sid_short = ""
        file_name = today_folder +"/" + sid_short + '-' + no_accent_vietnamese(obj.name.replace(" ",'')) + '-' +obj.lab_id+'_GENOTYPE.docx'
        pic_file = settings.STATIC_DIR + "/run_img/" + obj.lab_id + '_genotype.png'
        if obj.sid == None:
            sid = ""
        else:
            sid = str(obj.sid)
        
        if obj.age != None:
            age = str(obj.age)
        else:
            age = ""
        #start merging
        f_1 = MailMerge(settings.STATIC_DIR + '/word_template/HCVgeno FORM.docx')
        f_1.merge(
            sid = sid, 
            lab_id = str(obj.lab_id), 
            sex = obj.sex, 
            doctor = obj.doctor, 
            address = obj.address, 
            age = age, 
            name = obj.name, 
            date_receive = obj.added.strftime('%d/%m/%Y'), 
            date_finished = obj.modified.strftime('%d/%m/%Y-%H:%M'), 
            d= datetime.now().strftime("%d"), 
            m= datetime.now().strftime('%m'), 
            y= datetime.now().strftime('%Y'),
            genotype = obj_geno.genotype)
     
        f_1.write(file_name)

        doc = Document(file_name)
        for paragraph in doc.paragraphs:
            if r"None" in paragraph.text:
                print('YES NONE IS HERE')
                paragraph.text = paragraph.text.strip().replace("None", "")
            if "[image]" in paragraph.text:
                paragraph.text = paragraph.text.strip().replace("[image]", "")

                run = paragraph.add_run()
                run.add_picture(pic_file, width=Cm(12))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(file_name)
        completed = obj.lab_id + " vào " + today_folder
    except FileNotFoundError:
        incompleted = "Chưa có file hình " + obj.lab_id
        doc.save(file_name)
        completed = obj.lab_id + " vào " + today_folder
    except:
        incompleted = obj.lab_id
    return completed, incompleted
